#!/usr/bin/env python3
"""Update EIC Accelerator Short Proposal with strategic improvements."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import os

SRC = os.path.join(os.path.dirname(__file__), "EIC_Accelerator_Short_Proposal_Summary.docx")
DST = os.path.join(os.path.dirname(__file__), "EIC_Accelerator_Short_Proposal_Summary_v2.docx")

doc = Document(SRC)

# ── Helper: replace paragraph text preserving first-run formatting ──────────
def replace_para_text(para, new_text):
    """Replace all text in a paragraph, keeping the style of the first run."""
    if para.runs:
        fmt = para.runs[0].font
        font_name = fmt.name
        font_size = fmt.size
        font_bold = fmt.bold
        font_italic = fmt.italic
        font_color = fmt.color.rgb if fmt.color and fmt.color.rgb else None
    else:
        font_name = font_size = font_bold = font_italic = font_color = None

    # Clear existing runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold


def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = copy.deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    # Clear copied text
    for run in new_para.runs:
        run.text = ""
    if new_para.runs:
        new_para.runs[0].text = text
    else:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


# ── 1. EXECUTIVE SUMMARY – Strengthen with traction evidence ────────────────
# Para 5 – Main exec summary paragraph
doc.paragraphs[5].text = (
    "Lupe Analytics is building a context agent for marketing analytics: an AI system that "
    "deeply understands raw data sources, domain knowledge, business logic, and the specific "
    "context of each company and user. Unlike general-purpose AI analytics tools that require "
    "extensive manual setup and a dedicated data team, Lupe delivers enterprise-grade analytics "
    "intelligence to small and mid-sized mobile studios from day one. "
    "Lupe is not a concept — it is a production system with paying customers, a live AI backend "
    "processing real marketing data daily, and a test suite of over 3,000 automated tests "
    "validating every component of the platform."
)

# Para 6 – OpenAI reference (already strong, add one strategic sentence at end)
doc.paragraphs[6].text = (
    "The core innovation is a multi-layered context system for data agents. In February 2026, "
    "OpenAI published details of their in-house data agent: a system that reasons over 600+ "
    "petabytes and 70,000 datasets by combining schema metadata, human annotations, code-level "
    "enrichment, institutional knowledge, learning memory, and live queries into a unified "
    "context layer. OpenAI built this for themselves. It is not a product, and they have no "
    "plans to open-source it. Lupe is building the productized, generalizable version of this "
    "architecture — and has already deployed three of the six context layers in production."
)

# ── 2. SECTION 2.3 – Make Context Agent more concrete ───────────────────────
# Para 19: Strengthen the core innovation statement
doc.paragraphs[19].text = (
    "Lupe is building the productized, multi-tenant, vertical-specific version of this context "
    "agent architecture. The fundamental research challenge is: how do you build a context layer "
    "that works across hundreds of different companies, each with their own data sources, domain "
    "definitions, and business logic, and that keeps itself current as all three change "
    "continuously? Lupe's answer is the Context Agent — an autonomous AI system that "
    "continuously monitors, maintains, and evolves the semantic data layer for each client "
    "without human intervention."
)

# ── 2.5 – Add new subsection about the Context Agent (the R&D focus) ────────
# Find paragraph 30 (after the three capabilities) and add content after it
# Para 30 is the last capability (Ads Co-Pilot)
# We'll modify para 30 to add a bridge to the new content

# After the three capabilities (paras 28-30), before Section 2.4
# Para 30 text ends with "...the business context behind them."
# We want to add a paragraph about the Context Agent as the R&D core

# Find the right spot - after para 30 (last capability), before 2.4
p30 = doc.paragraphs[30]

# Add the Context Agent R&D focus paragraph
new_text_1 = (
    "The central R&D challenge — and the primary focus of this EIC funding request — is the "
    "Context Agent itself: an autonomous system that sits between a company's raw data "
    "infrastructure and its AI-powered analytics. The Context Agent automatically discovers "
    "data schemas, infers semantic relationships, builds and maintains a semantic layer "
    "(the structured mapping of business concepts to database columns), detects when schemas "
    "or business logic change, and self-repairs its understanding without human intervention. "
    "This is architecturally analogous to what Bridge — Lupe's internal AI development "
    "orchestrator — already does for code: it analyzes context, plans actions, executes them, "
    "and verifies results in an autonomous loop. The Context Agent applies this same pattern "
    "to data: analyze the data landscape, plan semantic mappings, execute enrichment, and "
    "verify accuracy — continuously and autonomously."
)

new_text_2 = (
    "Once the Context Agent maintains an accurate, up-to-date semantic layer for a client, "
    "it unlocks a second capability: the ability to spin up specialized, domain-aware AI agents "
    "on demand. A reporting agent that generates daily executive summaries. A QA agent that "
    "answers natural language questions with SQL-grade precision. An alerting agent that "
    "monitors KPIs and flags anomalies. A creative analytics agent that correlates ad creative "
    "variations with performance outcomes. Each agent inherits the full context stack from the "
    "Context Agent, eliminating the months of setup that enterprises currently spend configuring "
    "each analytics tool individually. Lupe has already built and deployed two such agents — "
    "the Reporting Agent and the QA Agent — in production, validating the architecture."
)

p31 = insert_paragraph_after(p30, new_text_1)
p32 = insert_paragraph_after(p31, new_text_2)

# ── 3. SECTION 2.4 – Strengthen TRL evidence ────────────────────────────────
# After inserting 2 paragraphs, indices shifted by 2
# Original para 32 is now at index 34
# Find the TRL paragraph by content
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Lupe is a working product with paying customers"):
        trl_idx = i
        break

doc.paragraphs[trl_idx].text = (
    "Lupe is a working product with paying customers. The current platform (TRL 6–7) "
    "integrates with major ad networks (Meta, Google, TikTok, AppLovin), MMPs (Adjust, "
    "AppsFlyer), and analytics platforms. The production backend includes: a QA Agent that "
    "translates natural language questions into semantic layer queries with sub-second latency; "
    "a Reporting Agent that generates structured daily and weekly executive summaries with "
    "evidence-backed insights; a Client Fast Path query compiler that bypasses the semantic "
    "layer for high-frequency queries, reducing latency by 80%; and a full-stack console "
    "with dynamic dashboards, KPI cards, chat interface, and collection pinning. The platform "
    "runs on a hardened AI pipeline with 3,000+ automated tests, strict model isolation "
    "between agents, and production-grade security. Layers 1 (Schema & Usage Context), "
    "2 (Domain Expert Annotations), and 4 (Organizational Context) of the six-layer context "
    "architecture are operational. Layers 3, 5, and 6 — the self-updating, learning, and "
    "code-enrichment layers that constitute the Context Agent's autonomous capabilities — "
    "are in active R&D and represent the core innovation this grant would accelerate."
)

# ── 4. COMPANY STATUS (Section 4.3) – Fill empty paragraph ──────────────────
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "" and i > 0:
        prev = doc.paragraphs[i-1]
        if "Company Status" in prev.text:
            p.text = (
                "Lupe Analytics UG is incorporated in Berlin, Germany (2024). The company is "
                "co-founder-led with a lean, execution-focused team. The product is live and in "
                "production with paying customers. Current MRR is €4,000 from a completed pilot "
                "engagement, with a qualified pipeline of European mobile gaming studios in active "
                "discussions. The founders have bootstrapped the company to product-market fit "
                "without external funding, demonstrating capital efficiency and technical velocity "
                "that de-risks the EIC investment."
            )
            break

# ── 5. HIRING PLAN (Section 4.4) – Strengthen ──────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("With EIC funding, Lupe plans to expand"):
        p.text = (
            "With EIC funding, Lupe plans to scale the team from 2 to 5 with three strategic "
            "hires. First, a Senior AI/ML Engineer to lead the Context Agent's self-updating "
            "mechanism — specifically semantic drift detection, automated schema re-enrichment, "
            "and the learning memory system. This role requires deep experience in knowledge "
            "representation and applied AI systems. Second, a Data Engineer to scale the "
            "multi-tenant platform architecture, optimize the query compilation pipeline, and "
            "extend warehouse support beyond BigQuery to Snowflake and Databricks. Third, a "
            "Go-to-Market Lead to drive customer acquisition across European mobile studios, "
            "manage the expansion into adjacent verticals, and build the partner ecosystem. "
            "All three hires will be based in Berlin, reinforcing Lupe's European talent base."
        )
        break

# ── 6. WHY EIC (Section 6) – Apply Voss + 48 Laws ──────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("The context agent for marketing analytics represents"):
        # Rewrite with tactical empathy + European sovereignty framing
        p.text = (
            "The EIC exists to ensure Europe builds and controls its own deep-tech future. "
            "The context agent for marketing analytics is precisely the kind of foundational "
            "technology where European leadership matters — and where it is most at risk. "
            "OpenAI has demonstrated that a six-layer context architecture is the key to making "
            "AI genuinely useful for enterprise data work. But they built it for themselves. "
            "The companies best positioned to productize this are US-based AI labs and "
            "well-funded Silicon Valley startups. Without European investment in this "
            "foundational technology, every European enterprise that wants AI-powered analytics "
            "will depend on US-controlled infrastructure processing their most sensitive "
            "business data and competitive intelligence."
        )
        break

for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Lupe is built in Berlin"):
        p.text = (
            "Lupe is built in Berlin by a team with a proven track record of building data "
            "infrastructure that drives real business outcomes — including the data platform "
            "behind a successful acquisition by a US public company. The founding team has "
            "lived the problem they are solving for six years. They are not entering this "
            "market with a thesis; they are entering it with production code, paying customers, "
            "and domain expertise that no horizontal AI tool can replicate. EIC support would "
            "enable Lupe to maintain a European technology lead in context-aware AI agents, "
            "scale commercially before US competitors close the window, and validate a context "
            "agent framework with applications far beyond marketing — in any vertical where "
            "business context is the bottleneck between raw data and decisions."
        )
        break

for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("The timing is critical"):
        p.text = (
            "The timing is critical and the window is narrow. Plurio's $3.5M seed round "
            "(February 2026) validates venture capital interest in this category. OpenAI's "
            "publication validates the architecture. Anthropic, Google, and OpenAI are all "
            "investing in agentic AI systems that will inevitably move into enterprise data. "
            "The question is not whether context agents will become the standard for business "
            "intelligence — it is whether Europe will build them or buy them. Lupe represents "
            "the opportunity for Europe to own the answer."
        )
        break

# ── 7. COMPETITIVE TABLE – Strengthen "Why Lupe Wins" column ────────────────
comp_table = doc.tables[0]
wins = [
    "No blending with product data; no AI interpretation layer; need a data team to connect the dots",
    "No self-updating context layer; no code-level enrichment; cannot maintain semantic models autonomously",
    "No ad spend allocation, no attribution intelligence, no cross-platform marketing context",
    "Require months of setup by a data team; no domain knowledge; no autonomous context maintenance",
    "Horizontal tools with zero domain expertise; no vertical context layers; no self-updating semantic models",
]
for row_idx, win_text in enumerate(wins):
    comp_table.rows[row_idx + 1].cells[2].text = win_text

# ── 8. KEY MILESTONES TABLE – Make more specific and ambitious ──────────────
milestone_table = doc.tables[2]
milestones = [
    ("M1–M6", "Context Agent v1 deployed: automated schema discovery and semantic layer generation for new clients. Multi-tenant platform live with automated onboarding. 10+ paying customers, €30K MRR. Patent filing for self-updating semantic layer architecture."),
    ("M7–M12", "Context Agent self-updating mechanism validated: semantic drift detection achieving >95% accuracy, automated re-enrichment operational. Learning memory (Layer 5) in production. Creative analytics and A/B testing modules launched. 25+ customers across 3+ European markets."),
    ("M13–M18", "First non-mobile vertical pilot (e-commerce marketing analytics). Custom agent spinning capability live — clients can configure domain-specific agents without engineering support. Automated campaign execution across all major ad networks. 50+ customers."),
    ("M19–M24", "Context Agent framework validated across 2+ verticals with documented generalization methodology. Full code-level enrichment (Layer 3) operational. Platform fully self-serve. TRL 9 achieved. Series A readiness with €500K+ ARR trajectory."),
]
for row_idx, (timeline, milestone) in enumerate(milestones):
    milestone_table.rows[row_idx + 1].cells[0].text = timeline
    milestone_table.rows[row_idx + 1].cells[1].text = milestone

# ── 9. FUNDING REQUEST – Strengthen use of proceeds ─────────────────────────
for i, p in enumerate(doc.paragraphs):
    if "R&D (50%)" in p.text:
        p.text = (
            "R&D (50% — €1.25M): Build the Context Agent's autonomous capabilities — "
            "self-updating semantic layer maintenance (Layer 6), code-level enrichment of "
            "data pipelines (Layer 3), learning memory from user corrections (Layer 5), and "
            "semantic drift detection with automated validation. Extend multi-tenant "
            "architecture for 100+ concurrent clients. Develop the custom agent spinning "
            "framework enabling per-client specialized agents."
        )
        break

for i, p in enumerate(doc.paragraphs):
    if "Team Expansion (30%)" in p.text:
        p.text = (
            "Team Expansion (30% — €750K): Hire Senior AI/ML Engineer (Context Agent lead), "
            "Data Engineer (multi-tenant scaling), and Go-to-Market Lead (European expansion). "
            "Scale team from 2 to 5. All hires Berlin-based."
        )
        break

for i, p in enumerate(doc.paragraphs):
    if "Go-to-Market (20%)" in p.text:
        p.text = (
            "Go-to-Market (20% — €500K): Customer acquisition through direct outreach to "
            "European mobile studios, presence at key industry events (GDC, PGC, Gamescom), "
            "and LinkedIn thought leadership. Validate context agent framework in 2–3 adjacent "
            "verticals (e-commerce, SaaS growth analytics). Build strategic partnerships with "
            "data platform providers (dbt Labs, Snowflake, BigQuery)."
        )
        break

# ── 10. PROBLEM SECTION – Add urgency ───────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Getting initial dashboards is easy"):
        p.text = (
            "Getting initial dashboards is easy. Unlocking the full value of this data requires "
            "bringing it in-house, building a unified data model, and hiring data engineers, "
            "analytics engineers, and analysts. Teams spend months setting up infrastructure "
            "and still end up with slow, disconnected insights while costs escalate at scale. "
            "Large studios build custom in-house solutions costing €500K+ annually in personnel "
            "alone. Small and mid-sized studios — which make up 80% of the market — compete "
            "against them but lack the resources to do the same. The result is a structural "
            "competitive disadvantage that no existing tool category addresses."
        )
        break

# ── SAVE ────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"✅ Saved updated proposal to: {DST}")
